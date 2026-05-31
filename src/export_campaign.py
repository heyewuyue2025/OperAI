"""Campaign Package 全景导出：9 章节结构化 Markdown / Word。

将任意 DAG 的 Agent 输出 + Verify Gate 结果 + 运行轨迹
组装为 PRD §7.2 定义的统一运营战役包。
"""
from __future__ import annotations

from typing import Any

from src.harness.verify_gate import VerifyResult

# ---------------------------------------------------------------------------
# 章节渲染注册表 — 每个 Agent 一个渲染函数
# ---------------------------------------------------------------------------

AGENT_LABELS: dict[str, str] = {
    "D": "数据洞察与证据",
    "U": "用户/客群策略",
    "C": "内容与标题",
    "A": "活动运营方案",
    "P": "产品运营洞察",
    "M": "市场定位与竞品",
    "F": "流量与预算分配",
    "N": "渠道与排期",
    "S": "社群运营策略",
    "E": "交易与转化",
}


def _b(item: str) -> str:
    """Markdown 列表项，跳过空字符串。"""
    return f"- {item}" if item else ""


def _bullets(items: list[str] | None, *, empty: str = "_（无）_") -> str:
    if not items:
        return empty
    return "\n".join(f"- {x}" for x in items)


def _fallback_block(agent_outputs: dict[str, dict[str, Any]]) -> str:
    lines: list[str] = []
    for aid, blob in agent_outputs.items():
        if isinstance(blob, dict) and blob.get("_operai_fallback"):
            label = AGENT_LABELS.get(aid, aid)
            lines.append(f"- **{label}（{aid}-Agent）**：{blob['_operai_fallback']}")
    if not lines:
        return ""
    return "\n> ⚠ **系统提示（降级输出）**\n>\n" + "\n".join(f"> {line}" for line in lines) + "\n"


# ---- 各 Agent 章节渲染 ----

def _render_d(payload: dict[str, Any]) -> str:
    insights = payload.get("insights") or []
    angles = payload.get("angles") or []
    risks = payload.get("risk_flags") or []
    evidence = payload.get("evidence_spans") or []

    ev_lines = []
    for e in evidence:
        if isinstance(e, dict):
            ev_lines.append(f"- **{e.get('field', '')}**：`{e.get('snippet', '')}`")
    ev_md = "\n".join(ev_lines) if ev_lines else "_（无摘录）_"

    return f"""### 关键判断

{_bullets([str(x) for x in insights])}

### 传播角度

{_bullets([str(x) for x in angles])}

### 风险与合规提示

{_bullets([str(x) for x in risks], empty="_（无风险标记）_")}

### 证据摘录（可溯源）

{ev_md}
"""


def _render_u(payload: dict[str, Any]) -> str:
    segments = payload.get("segments") or []
    lifecycle = payload.get("lifecycle_stage", "_未识别_")
    actions = payload.get("retention_actions") or []
    churn = payload.get("churn_risks") or []

    seg_md = ""
    for s in segments:
        if isinstance(s, dict):
            prio = s.get("priority", "")
            prio_label = {"high": "[高]", "medium": "[中]", "low": "[低]"}.get(prio, "")
            seg_md += f"- {prio_label} **{s.get('name', '')}**：{s.get('description', '')}\n"

    act_md = ""
    for a in actions:
        if isinstance(a, dict):
            act_md += f"- **{a.get('segment', '')}** → {a.get('action', '')}（渠道：{a.get('channel', '')}）\n"

    return f"""### 用户分群

{seg_md if seg_md else "_（无）_"}

### 生命周期阶段

`{lifecycle}`

### 留存/触达动作

{act_md if act_md else "_（无）_"}

### 流失风险

{_bullets([str(x) for x in churn])}
"""


def _render_c(payload: dict[str, Any], *, draft_overrides: dict[str, str] | None = None) -> str:
    drafts = dict(payload.get("drafts") or {})
    if draft_overrides:
        drafts.update({k: v for k, v in draft_overrides.items() if v is not None})

    titles = payload.get("title_variants") or []
    script = (payload.get("short_video_script") or "").strip()
    compliance = payload.get("compliance_notes") or []

    labels = {"weibo": "微博", "wechat": "公众号", "xhs": "小红书"}

    draft_md = ""
    for plat, text in drafts.items():
        draft_md += f"#### {labels.get(plat, plat)}\n\n{str(text).strip()}\n\n"

    return f"""### 备选标题

{_bullets([f"{i+1}. {t}" for i, t in enumerate(titles)]) if titles else "_（无）_"}

### 短视频口播稿

{script if script else "_（无）_"}

### 分平台文案

{draft_md if draft_md else "_（无文案）_"}

### 合规说明

{_bullets([str(x) for x in compliance]) if compliance else "_（无）_"}
"""


def _render_a(payload: dict[str, Any]) -> str:
    plan = payload.get("campaign_plan") or []
    budget = payload.get("budget_hints") or []
    roi = payload.get("roi_estimate") or {}

    plan_md = ""
    for p in plan:
        if isinstance(p, dict):
            tasks = "\n".join(f"  - {t}" for t in (p.get("tasks") or []))
            plan_md += f"- **{p.get('phase', '')}**（目标：{p.get('objective', '')}）\n{tasks}\n"

    budget_md = ""
    for b in budget:
        if isinstance(b, dict):
            budget_md += f"- {b.get('channel', '')}：{b.get('suggestion', '')}"
            if b.get("percent_range"):
                budget_md += f"（建议占比 {b['percent_range']}）"
            budget_md += "\n"

    return f"""### 战役计划

{plan_md if plan_md else "_（无）_"}

### 预算建议

{budget_md if budget_md else "_（无）_"}

### ROI 预估

{roi.get('summary', '_（无）_') if isinstance(roi, dict) else '_（无）_'}

### 假设与置信度

{_bullets(roi.get('assumptions', [])) if isinstance(roi, dict) and roi.get('assumptions') else '_（无）_'}

置信度：{roi.get('confidence', '_未评估_') if isinstance(roi, dict) else '_未评估_'}
"""


def _render_p(payload: dict[str, Any]) -> str:
    features = payload.get("feature_insights") or []
    ux = payload.get("ux_signals") or []
    iterations = payload.get("iteration_hints") or []

    feat_md = ""
    for f in features:
        if isinstance(f, dict):
            feat_md += f"- **{f.get('feature', '')}**：{f.get('insight', '')}\n"

    iter_md = ""
    for i in iterations:
        if isinstance(i, dict):
            iter_md += f"- [{i.get('priority', '')}] {i.get('recommendation', '')}\n"

    return f"""### 功能洞察

{feat_md if feat_md else "_（无）_"}

### 用户体验信号

{_bullets([str(x) for x in ux])}

### 迭代建议

{iter_md if iter_md else "_（无）_"}
"""


def _render_m(payload: dict[str, Any]) -> str:
    positioning = payload.get("positioning", "")
    competitive = payload.get("competitive_notes") or []
    channel_mix = payload.get("channel_mix") or []

    comp_md = ""
    for c in competitive:
        if isinstance(c, dict):
            comp_md += f"- **{c.get('topic', '')}**：{c.get('note', '')}\n"

    ch_md = ""
    for ch in channel_mix:
        if isinstance(ch, dict):
            ch_md += f"- {ch.get('channel', '')}（角色：{ch.get('role', '')}，权重：{ch.get('weight', '')}）\n"

    return f"""### 定位

{positioning if positioning else '_（未定义）_'}

### 竞品观察

{comp_md if comp_md else "_（无）_"}

### 渠道组合

{ch_md if ch_md else "_（无）_"}
"""


def _render_f(payload: dict[str, Any]) -> str:
    scores = payload.get("channel_scores") or []
    budget = payload.get("budget_allocation") or []
    hints = payload.get("conversion_hints") or []

    scores_md = ""
    for s in scores:
        if isinstance(s, dict):
            scores_md += f"- **{s.get('channel', '')}**：{s.get('score', '—')} 分 — {s.get('rationale', '')}\n"

    budget_md = ""
    for b in budget:
        if isinstance(b, dict):
            budget_md += f"- {b.get('channel', '')}：{b.get('percent', '—')}%\n"

    return f"""### 渠道评分

{scores_md if scores_md else "_（无）_"}

### 预算分配建议

{budget_md if budget_md else "_（无）_"}

### 转化优化提示

{_bullets([str(x) for x in hints])}
"""


def _render_n(payload: dict[str, Any]) -> str:
    sched = payload.get("schedule_suggestions") or []
    tags = payload.get("hashtags") or []
    notes = payload.get("platform_notes") or {}

    sched_md = ""
    if sched:
        sched_md = "| 平台 | 时间窗 | 理由 |\n|------|--------|------|\n"
        for s in sched:
            if isinstance(s, dict):
                sched_md += f"| {s.get('platform', '')} | {s.get('window', '')} | {s.get('reason', '')} |\n"
    else:
        sched_md = "_（无）_"

    notes_md = ""
    if isinstance(notes, dict) and notes:
        notes_md = "\n".join(f"- **{k}**：{v}" for k, v in notes.items())
    else:
        notes_md = "_（无）_"

    first_comments = payload.get("first_comment_suggestions") or []
    fc_md = ""
    for fc in first_comments:
        if isinstance(fc, dict):
            fc_md += f"- **{fc.get('platform', '')}**：{fc.get('text', '')}\n"

    return f"""### 排期建议

{sched_md}

### 话题标签

{_bullets([str(x) for x in tags])}

### 分平台注意事项

{notes_md}

### 首评引导建议

{fc_md if fc_md else "_（无）_"}
"""


def _render_s(payload: dict[str, Any]) -> str:
    actions = payload.get("community_actions") or []
    kol = payload.get("kol_hints") or []
    scripts = payload.get("engagement_scripts") or []

    act_md = ""
    for a in actions:
        if isinstance(a, dict):
            act_md += f"- **{a.get('action', '')}**（时机：{a.get('timing', '')}，目标：{a.get('goal', '')}）\n"

    kol_md = ""
    for k in kol:
        if isinstance(k, dict):
            kol_md += f"- {k.get('profile', '')}：{k.get('approach', '')}\n"

    scr_md = ""
    for s in scripts:
        if isinstance(s, dict):
            scr_md += f"- **{s.get('scenario', '')}**：{s.get('script', '')}\n"

    return f"""### 社群动作

{act_md if act_md else "_（无）_"}

### KOL 策略

{kol_md if kol_md else "_（无）_"}

### 互动话术

{scr_md if scr_md else "_（无）_"}
"""


def _render_e(payload: dict[str, Any]) -> str:
    funnel = payload.get("funnel_steps") or []
    promos = payload.get("promo_suggestions") or []
    ctas = payload.get("cta_variants") or []

    funnel_md = ""
    for f in funnel:
        if isinstance(f, dict):
            drop = f"（流失风险：{f.get('dropoff_risk', '')}）" if f.get("dropoff_risk") else ""
            funnel_md += f"- **{f.get('step', '')}**：{f.get('description', '')}{drop}\n"

    promo_md = ""
    for p in promos:
        if isinstance(p, dict):
            promo_md += f"- {p.get('offer', '')}（约束：{p.get('constraint', '')}）\n"

    return f"""### 转化漏斗

{funnel_md if funnel_md else "_（无）_"}

### 促销建议

{promo_md if promo_md else "_（无）_"}

### CTA 变体

{_bullets([str(x) for x in ctas])}
"""


_CHAPTER_RENDERERS: dict[str, Any] = {
    "D": _render_d,
    "U": _render_u,
    "C": _render_c,
    "A": _render_a,
    "P": _render_p,
    "M": _render_m,
    "F": _render_f,
    "N": _render_n,
    "S": _render_s,
    "E": _render_e,
}

# ---------------------------------------------------------------------------
# 主构建函数
# ---------------------------------------------------------------------------


def build_campaign_markdown(
    *,
    title: str,
    task_id: str,
    run_id: str,
    pack_id: str,
    dag: list[str],
    agent_outputs: dict[str, dict[str, Any]],
    verify_result: VerifyResult | None = None,
    trace_steps: list[dict[str, Any]] | None = None,
    draft_overrides: dict[str, str] | None = None,
) -> str:
    """生成 9 章节 Campaign Package Markdown。

    动态章节由 DAG 决定：DAG 中每个 Agent 对应一个业务章节；
    「合规与审核记录」与「运行轨迹附录」固定在末尾。
    """
    lines: list[str] = []
    lines.append(f"# OperAI 运营战役包")
    lines.append("")
    lines.append(f"> **任务**：{title}")
    lines.append(f"> **编排链路**：{' → '.join(dag)}")
    lines.append(f"> **task_id**：`{task_id}` · **run_id**：`{run_id}`")
    lines.append("")

    fb = _fallback_block(agent_outputs)
    if fb:
        lines.append(fb)
        lines.append("")

    lines.append("---")
    lines.append("")

    # 章节编号从 1 开始
    chapter_idx = 1

    for agent_id in dag:
        aid = agent_id.upper()
        payload = agent_outputs.get(aid) or {}
        label = AGENT_LABELS.get(aid, aid)

        lines.append(f"## {chapter_idx}. {label}（{aid}-Agent）")
        lines.append("")

        renderer = _CHAPTER_RENDERERS.get(aid)
        if renderer:
            if aid == "C":
                lines.append(renderer(payload, draft_overrides=draft_overrides))
            else:
                lines.append(renderer(payload))
        else:
            lines.append(f"_Agent {aid} 无渲染器_")
            lines.append("")

        chapter_idx += 1

    # ---- 合规与审核记录 ----
    lines.append(f"## {chapter_idx}. 合规与审核记录")
    lines.append("")
    chapter_idx += 1

    if verify_result:
        vr = verify_result
        if vr.block_export:
            lines.append("> ⛔ **导出门禁：已触发** — 以下问题阻断导出，需人工处理后方可交付。")
        elif vr.warnings:
            lines.append("> 🟡 **导出门禁：通过（含提示）** — 建议人工复核以下项目。")
        else:
            lines.append("> ✅ **导出门禁：通过** — 未发现阻断项。")
        lines.append("")

        for check in vr.checks:
            icon = "✅" if check.get("passed") else "❌"
            lines.append(f"- {icon} **{check.get('id', '')}**：{check.get('detail', '')}")

        if vr.warnings:
            lines.append("")
            lines.append("### 审核提示")
            for w in vr.warnings:
                lines.append(f"- {w}")
    else:
        lines.append("_未执行验证。_")

    lines.append("")
    lines.append("### 人工审核状态")
    lines.append("")
    lines.append("- [ ] 已核对事实与数据来源")
    lines.append("- [ ] 已处理风险标记与敏感表述")
    lines.append("- [ ] 已确认平台规则、发布时间和风险提示")
    lines.append("")

    # ---- 运行轨迹附录 ----
    lines.append(f"## {chapter_idx}. 运行轨迹附录")
    lines.append("")
    chapter_idx += 1

    if trace_steps:
        lines.append("| 步骤 | 状态 | 耗时(ms) | 降级 | 输出摘要 |")
        lines.append("|------|------|----------|------|----------|")
        for step in trace_steps:
            status = step.get("status", "—")
            icon = {"success": "✅", "failed": "❌", "running": "⏳"}.get(status, "⚪")
            fallback_flag = "⚠" if step.get("fallback_flag") else ""
            snippet = str(step.get("snippet", ""))[:80]
            lines.append(
                f"| {step.get('step', '—')} | {icon} {status} | {step.get('duration_ms', '—')} | {fallback_flag} | {snippet} |"
            )
    else:
        lines.append("_无轨迹数据。_")

    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(
        f"*本战役包由 OperAI Harness 自动生成（链路: {' → '.join(dag)}）。发布前请人工复核全部内容。*"
    )

    return "\n".join(lines)


def build_campaign_docx_bytes(
    *,
    title: str,
    task_id: str,
    run_id: str,
    pack_id: str,
    dag: list[str],
    agent_outputs: dict[str, dict[str, Any]],
    verify_result: VerifyResult | None = None,
    trace_steps: list[dict[str, Any]] | None = None,
    draft_overrides: dict[str, str] | None = None,
) -> bytes:
    """生成 Campaign Package Word（.docx）字节。

    结构与 Markdown 版一致，9 章节动态渲染。
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError as e:
        raise RuntimeError("请安装 python-docx：pip install python-docx") from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    doc.add_heading("OperAI 运营战役包", level=0)
    doc.add_paragraph(f"任务：{title}")
    doc.add_paragraph(f"编排链路：{' → '.join(dag)}")
    doc.add_paragraph(f"task_id: {task_id}")
    doc.add_paragraph(f"run_id: {run_id}")

    # 降级警示
    for aid, blob in agent_outputs.items():
        if isinstance(blob, dict) and blob.get("_operai_fallback"):
            label = AGENT_LABELS.get(aid, aid)
            doc.add_paragraph(f"[降级] {label}（{aid}-Agent）: {blob['_operai_fallback']}")

    chapter_idx = 1

    for agent_id in dag:
        aid = agent_id.upper()
        payload = agent_outputs.get(aid) or {}
        label = AGENT_LABELS.get(aid, aid)

        doc.add_heading(f"{chapter_idx}. {label}（{aid}-Agent）", level=1)
        chapter_idx += 1

        # 为每个 Agent 将 Markdown 渲染结果转换为 Word 段落
        renderer = _CHAPTER_RENDERERS.get(aid)
        if renderer:
            if aid == "C":
                md = renderer(payload, draft_overrides=draft_overrides)
            else:
                md = renderer(payload)
            for block in md.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("### "):
                    doc.add_heading(block[4:], level=2)
                elif block.startswith("#### "):
                    doc.add_heading(block[5:], level=3)
                elif block.startswith("| "):
                    # 表格
                    rows = [r for r in block.split("\n") if r.startswith("|")]
                    if len(rows) >= 2:
                        cells = [c.strip() for c in rows[0].split("|")[1:-1]]
                        table = doc.add_table(rows=len(rows), cols=len(cells))
                        table.style = "Light Grid Accent 1"
                        for ri, row in enumerate(rows):
                            for ci, cell in enumerate(c.strip().split("|")[1:-1]):
                                table.cell(ri, ci).text = cell.strip()
                elif block.startswith("- "):
                    for item in block.split("\n"):
                        item = item.strip()
                        if item.startswith("- "):
                            doc.add_paragraph(item[2:], style="List Bullet")
                else:
                    doc.add_paragraph(block)

    # ---- 合规与审核记录 ----
    doc.add_heading(f"{chapter_idx}. 合规与审核记录", level=1)
    chapter_idx += 1

    if verify_result:
        vr = verify_result
        if vr.block_export:
            doc.add_paragraph("⛔ 导出门禁：已触发 — 以下问题阻断导出，需人工处理。")
        elif vr.warnings:
            doc.add_paragraph("🟡 导出门禁：通过（含提示）— 建议人工复核。")
        else:
            doc.add_paragraph("✅ 导出门禁：通过 — 未发现阻断项。")

        for check in vr.checks:
            icon = "✓" if check.get("passed") else "✗"
            doc.add_paragraph(f"{icon} {check.get('id', '')}: {check.get('detail', '')}")

        if vr.warnings:
            doc.add_heading("审核提示", level=2)
            for w in vr.warnings:
                doc.add_paragraph(w, style="List Bullet")

    doc.add_heading("人工审核状态", level=2)
    doc.add_paragraph("☐ 已核对事实与数据来源")
    doc.add_paragraph("☐ 已处理风险标记与敏感表述")
    doc.add_paragraph("☐ 已确认平台规则、发布时间和风险提示")

    # ---- 运行轨迹附录 ----
    doc.add_heading(f"{chapter_idx}. 运行轨迹附录", level=1)
    chapter_idx += 1

    if trace_steps:
        table = doc.add_table(rows=1 + len(trace_steps), cols=4)
        table.style = "Light Grid Accent 1"
        for ci, header in enumerate(["步骤", "状态", "耗时(ms)", "输出摘要"]):
            table.cell(0, ci).text = header
        for ri, step in enumerate(trace_steps, 1):
            table.cell(ri, 0).text = step.get("step", "—")
            table.cell(ri, 1).text = str(step.get("status", "—"))
            table.cell(ri, 2).text = str(step.get("duration_ms", "—"))
            snippet = str(step.get("snippet", ""))[:120]
            if step.get("fallback_flag"):
                snippet = "[降级] " + snippet
            table.cell(ri, 3).text = snippet
    else:
        doc.add_paragraph("无轨迹数据。")

    doc.add_paragraph("")
    doc.add_paragraph(
        f"本战役包由 OperAI Harness 自动生成（链路: {' → '.join(dag)}）。发布前请人工复核全部内容。"
    )

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
